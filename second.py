import os
import time
import requests
import json
import pandas as pd
import openai
from langchain.docstore.document import Document
from langchain.chains.summarize import load_summarize_chain
# from langchain.llms.openai import OpenAI, OpenAIChat
from langchain_community.llms import OpenAI, OpenAIChat
from langchain_openai import AzureChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.text_splitter import CharacterTextSplitter
pd.set_option('display.max_rows', 100)




# Set endpoint URL
endpoint = "https://tuopenai.openai.azure.com" # TU endponint URL. O

# Set Authenication code
f = open("azureopenaikey.txt", "r")
api_key = f.read()
f.close()
deployment_id = "TUChatGPT"
api_version = "2024-02-01"





# Read the mannually processed data
df = pd.read_excel('20240419_c.xlsx')
print(df.shape)
# display(df.head())





# Set the contents of each agenda 
# Please mannually adjust the coding here to include every agenda discussed at the meeting as above
target_agenda = df['Agenda'].unique()[1]
agenda_1 = df[df['Agenda'] == target_agenda]['Speech'].str.cat(sep=' ')

target_agenda = df['Agenda'].unique()[2]
agenda_2 = df[df['Agenda'] == target_agenda]['Speech'].str.cat(sep=' ')

target_agenda = df['Agenda'].unique()[3]
agenda_3 = df[df['Agenda'] == target_agenda]['Speech'].str.cat(sep=' ')

target_agenda = df['Agenda'].unique()[4]
agenda_4 = df[df['Agenda'] == target_agenda]['Speech'].str.cat(sep=' ')






# Convert the text data agenda_x into langchain supoprted Document type data
from langchain.docstore.document import Document
from langchain.text_splitter import CharacterTextSplitter


text_splitter = CharacterTextSplitter(
    separator=".",
    chunk_size=3000,# number of text which openAI API can read
    chunk_overlap=200,# number of text to overlap for open API to understand the context
    length_function=len,

)
def create_docs_for_agenda(agenda_text):
    texts = text_splitter.split_text(agenda_text)
    docs = []
    for text in texts:
        doc = Document(page_content=text)
        docs.append(doc)
    print(f"Total documents created: {len(docs)}")  # デバッグ情報を出力
    return docs

# 各アジェンダについてドキュメントを生成し、生成されたドキュメントの数を確認
print(f"length of raw text data: {len(agenda_1)}")
docs_1 = create_docs_for_agenda(agenda_1)
print(f"Length of docs_1: {len(docs_1)}")





# Convert the text data agenda_x into langchain supoprted Document type data
chunk_size = 3000 # number of text which openAI API can read
chunk_overlap = 200 # number of text to overlap for open API to understand the context
text_splitter = CharacterTextSplitter(
    chunk_size=chunk_size,
    chunk_overlap=chunk_overlap,
    length_function=len,
)

# convert each text data (agend_x) into langchian data 
def create_docs_for_agenda(agenda_text):
    texts = text_splitter.split_text(agenda_text)
    docs = []
    for text in texts:
        doc = Document(page_content=text)
        docs.append(doc)
    return docs

# Please mannually adjust the coding here to include every agenda discussed at the meeting as above
docs_1 = create_docs_for_agenda(agenda_1)
docs_2 = create_docs_for_agenda(agenda_2)
docs_3 = create_docs_for_agenda(agenda_3)
docs_4 = create_docs_for_agenda(agenda_4)
#docs_5 = create_docs_for_agenda(agenda_5)
#docs_6 = create_docs_for_agenda(agenda_6)



# Set endpoint and authenticatio for langchain
os.environ["AZURE_OPENAI_API_KEY"] = api_key
os.environ["AZURE_OPENAI_ENDPOINT"] = endpoint





from docx import Document

# 新しいWordドキュメントを作成するか、既存のドキュメントを開く関数
def add_to_document(contents, doc_path="meeting_summary.docx"):
    try:
        # 既存のドキュメントを開く
        doc = Document(doc_path)
    except FileNotFoundError:
        # ファイルが存在しない場合、新しいドキュメントを作成
        doc = Document()

    # 出力内容を文に分割し、Wordドキュメントに追加
    sentences = contents.split('. ')
    for sentence in sentences:
        doc.add_paragraph(f"- {sentence.strip()}")

    # ドキュメントを保存
    doc.save(doc_path)


openaichat = AzureChatOpenAI(temperature=0, model="gpt-3.5-turbo", deployment_name=deployment_id, openai_api_version=api_version)

#VALUE CHAIN#####################################################

# Adjust the prompt according to the meeting contnets and speaker
# prompt_template = """As a CEO of the company, listen to the report provided at the status meeting by After sales general manager of a company in Ukraine.
# Company in located in Ukraine and is a official distributor of Toyota and Lexus vehicles and spare parts. 
# Please compile a minutes of the meeting report as you are a CEO and make it reopr to General director of the company based on the provided information.
# Your memo should adhere to the following guidelines:
#     - Tone: Formal
#     - Format: Detailed report divided into chapters as outlined below
#     - Length: Approximately 3000-4000 words
#     - Focus on the subject matter of discussions rather than the participants
#     - Provide detailed explanations and analysis for each topic
#     - Prepare report with bullet points for each chapter. Bullet moints are must. After each bullet poit make line break
#     - Please exlude introduction, greeting, closing part and signature. Only report
#     - Organize the memo into chapters on:
#         1. Service Unit
#         2. Spare parts performance KPI
#         3. Spare parts sales
#         4. parallel import
#         5. Used Car Sales
#         6. Profitability

# Detailed insights, data comparisons, market trends, and strategic conclusions are crucial. Where applicable, include follow-up actions with specific owners assigned.

# Your analysis should incorporate the following discussions and data points:
# {text}

# Ensure precision and depth in your report:"""
# PROMPT = PromptTemplate(template=prompt_template, input_variables=["text"])

# chain = load_summarize_chain(
#         openaichat,
#         chain_type="refine",  # This might need adjustment based on the actual functionality available.
#         return_intermediate_steps=True,
#         question_prompt=PROMPT
#     )

# resp = chain({"input_documents": docs_3}, return_only_outputs=True)

# #add_to_document(resp["output_text"])

# resp["output_text"]

# CONTENTS = resp["output_text"]

# # 文を分割し、箇条書きで表示
# sentences = CONTENTS.split('. ')
# for sentence in sentences:
#     print(f"- {sentence.strip()}")






#CUSTOMER FIRST#####################################################

# prompt_template = """As a CEO of the company, listen to the report provided at the status meeting by Technical and Warranty general manager a company.
# Company in located in Ukraine and is a official distributor of Toyota and Lexus vehicles and spare parts. 
# Please compile a minutes of the meeting report as you are a CEO and make it reopr to General director of the company based on the provided information.
# Your memo should adhere to the following guidelines:
#     - Tone: Formal
#     - Format: Detailed report divided into chapters as outlined below
#     - Length: Approximately 3000-4000 words
#     - Focus on the subject matter of discussions rather than the participants
#     - Provide detailed explanations and analysis for each topic
#     - Prepare report with bullet points for each chapter. Bullet moints are must. After each bullet poit make line break
#     - Please exlude introduction, greeting, closing part and signature. Only report


# Detailed insights, data comparisons, market trends, and strategic conclusions are crucial. Where applicable, include follow-up actions with specific owners assigned.

# Your analysis should incorporate the following discussions and data points:
# {text}

# Ensure precision and depth in your report:"""
# PROMPT = PromptTemplate(template=prompt_template, input_variables=["text"])

# chain = load_summarize_chain(
#         openaichat,
#         chain_type="refine",  # This might need adjustment based on the actual functionality available.
#         return_intermediate_steps=True,
#         question_prompt=PROMPT
#     )

# resp = chain({"input_documents": docs_3}, return_only_outputs=True)

# #add_to_document(resp["output_text"])

# resp["output_text"]

# CONTENTS = resp["output_text"]

# # 文を分割し、箇条書きで表示
# sentences = CONTENTS.split('. ')
# for sentence in sentences:
#     print(f"- {sentence.strip()}")





#SALES#####################################################

# Adjust the prompt according to the meeting contnets and speaker
prompt_template = """As a CEO of the company, listen to the report provided at the status meeting by Lexus Sale General manager of the comppany.
Company in located in Ukraine and is a official distributor of Toyota and Lexus vehicles and spare parts.
Please compile a minutes of the meeting report as you are a CEO and make it reopr to General director of the company based on the provided information.
Your memo should adhere to the following guidelines:
    - Tone: Formal
    - Format: Detailed report divided into chapters 
    - Length: Approximately 3000-4000 words
    - Focus on the subject matter of discussions rather than the participants
    - Provide detailed explanations and analysis for each topic
    - Prepare report with bullet points for each chapter. Bullet moints are must. After each bullet poit make line break
    - Please exlude introduction, greeting, closing part and signature. Only report
    
Detailed insights, data comparisons, market trends, and strategic conclusions are crucial. Where applicable, include follow-up actions with specific owners assigned.

Your analysis should incorporate the following discussions and data points:
{text}

Ensure precision and depth in your report:"""
PROMPT = PromptTemplate(template=prompt_template, input_variables=["text"])

chain = load_summarize_chain(
        openaichat,
        chain_type="refine",  # This might need adjustment based on the actual functionality available.
        return_intermediate_steps=True,
        question_prompt=PROMPT
    )

resp = chain({"input_documents": docs_2}, return_only_outputs=True)

#add_to_document(resp["output_text"])

resp["output_text"]

CONTENTS = resp["output_text"]

# 文を分割し、箇条書きで表示
sentences = CONTENTS.split('. ')
for sentence in sentences:
    print(f"- {sentence.strip()}")






#NEW BUSINESSES##################################    

# prompt_template = """As a CEO of the company, listen to the report provided at the status meeting by New Business Development general manager of a Company.
# Company in located in Ukraine and is a official distributor of Toyota and Lexus vehicles and spare parts. 
# Please compile a minutes of the meeting report as you are a CEO and make it reopr to General director of the company based on the provided information.
# Your memo should adhere to the following guidelines:
#     - Tone: Formal
#     - Format: Detailed report divided into chapters as outlined below
#     - Length: Approximately 3000-4000 words
#     - Focus on the subject matter of discussions rather than the participants
#     - Provide detailed explanations and analysis for each topic
#     - Prepare report with bullet points for each chapter. Bullet moints are must. After each bullet poit make line break
#     - Please exlude introduction, greeting, closing part and signature. Only report
#     - Organize the memo into chapters on:
#         1. Insurance project
#         2. Lease project for external customers
#         3. Lease for company employee
#         4. TFS communication
#         5. Others
        
# Detailed insights, data comparisons, market trends, and strategic conclusions are crucial. Where applicable, include follow-up actions with specific owners assigned.

# Your analysis should incorporate the following discussions and data points:
# {text}

# Ensure precision and depth in your report:"""
# PROMPT = PromptTemplate(template=prompt_template, input_variables=["text"])

# chain = load_summarize_chain(
#         openaichat,
#         chain_type="refine",  # This might need adjustment based on the actual functionality available.
#         return_intermediate_steps=True,
#         question_prompt=PROMPT
#     )

# resp = chain({"input_documents": docs_3}, return_only_outputs=True)

# #add_to_document(resp["output_text"])

# resp["output_text"]

# CONTENTS = resp["output_text"]

# # 文を分割し、箇条書きで表示
# sentences = CONTENTS.split('. ')
# for sentence in sentences:
#     print(f"- {sentence.strip()}")





#MARKETING##################################    

# prompt_template = """As a CEO of the company, listen to the report provided at the status meeting by Marketing general manager of a Company.
# Company in located in Ukraine and is a official distributor of Toyota and Lexus vehicles and spare parts. 
# Please compile a minutes of the meeting report as you are a CEO and make it reopr to General director of the company based on the provided information.
# Your memo should adhere to the following guidelines:
#     - Tone: Formal
#     - Format: Detailed report divided into chapters as outlined below
#     - Length: Approximately 3000-4000 words
#     - Focus on the subject matter of discussions rather than the participants
#     - Provide detailed explanations and analysis for each topic
#     - Prepare report with bullet points for each chapter. Bullet moints are must. After each bullet poit make line break
#     - Please exlude introduction, greeting, closing part and signature. Only report
#     - Organize the memo into chapters on:
#         1. New cars sales marketing
#         2. After sales and used cars sales marketing
#         3. Others

        
# Detailed insights, data comparisons, market trends, and strategic conclusions are crucial. Where applicable, include follow-up actions with specific owners assigned.

# Your analysis should incorporate the following discussions and data points:
# {text}

# Ensure precision and depth in your report:"""
# PROMPT = PromptTemplate(template=prompt_template, input_variables=["text"])

# chain = load_summarize_chain(
#         openaichat,
#         chain_type="refine",  # This might need adjustment based on the actual functionality available.
#         return_intermediate_steps=True,
#         question_prompt=PROMPT
#     )

# resp = chain({"input_documents": docs_3}, return_only_outputs=True)

# #add_to_document(resp["output_text"])

# resp["output_text"]

# CONTENTS = resp["output_text"]

# # 文を分割し、箇条書きで表示
# sentences = CONTENTS.split('. ')
# for sentence in sentences:
#     print(f"- {sentence.strip()}")






#FINANCE(TREASURY)##################################    

# prompt_template = """As a CEO of the company, listen to the report provided at the status meeting by Finance general manager of a Tcompany.
# Company in located in Ukraine and is a official distributor of Toyota and Lexus vehicles and spare parts. 
# Please compile a minutes of the meeting report as you are a CEO and make it reopr to General director of the company based on the provided information.
# Your memo should adhere to the following guidelines:
#     - Tone: Formal
#     - Format: Detailed report divided into chapters as outlined below
#     - Length: Approximately 3000-4000 words
#     - Focus on the subject matter of discussions rather than the participants
#     - Provide detailed explanations and analysis for each topic
#     - Prepare report with bullet points for each chapter. Bullet moints are must. After each bullet poit make line break
#     - Please exlude introduction, greeting, closing part and signature. Only report
#     - Organize the memo into chapters on:
#         1. Profit and loss 
#         2. Balance Sheet
#         3. Macro Economy
        
# Detailed insights, data comparisons, market trends, and strategic conclusions are crucial. Where applicable, include follow-up actions with specific owners assigned.

# Your analysis should incorporate the following discussions and data points:
# {text}

# Ensure precision and depth in your report:"""
# PROMPT = PromptTemplate(template=prompt_template, input_variables=["text"])

# chain = load_summarize_chain(
#         openaichat,
#         chain_type="refine",  # This might need adjustment based on the actual functionality available.
#         return_intermediate_steps=True,
#         question_prompt=PROMPT
#     )

# resp = chain({"input_documents": docs_3}, return_only_outputs=True)

# #add_to_document(resp["output_text"])

# resp["output_text"]

# CONTENTS = resp["output_text"]

# # 文を分割し、箇条書きで表示
# sentences = CONTENTS.split('. ')
# for sentence in sentences:
#     print(f"- {sentence.strip()}")






#GERERAL AFFAIRS##################################    

# prompt_template = """As a CEO of the company, listen to the report provided at the status meeting by General affairs general manager of a company.
# Company in located in Ukraine and is a official distributor of Toyota and Lexus vehicles and spare parts. 
# Please compile a minutes of the meeting report as you are a CEO and make it reopr to General director of the company based on the provided information.
# Your memo should adhere to the following guidelines:
#     - Tone: Formal
#     - Format: Detailed report divided into chapters as outlined below
#     - Length: Approximately 3000-4000 words
#     - Focus on the subject matter of discussions rather than the participants
#     - Provide detailed explanations and analysis for each topic
#     - Prepare report with bullet points for each chapter. Bullet moints are must. After each bullet poit make line break
#     - Please exlude introduction, greeting, closing part and signature. Only report

        
# Detailed insights, data comparisons, market trends, and strategic conclusions are crucial. Where applicable, include follow-up actions with specific owners assigned.

# Your analysis should incorporate the following discussions and data points:
# {text}

# Ensure precision and depth in your report:"""
# PROMPT = PromptTemplate(template=prompt_template, input_variables=["text"])

# chain = load_summarize_chain(
#         openaichat,
#         chain_type="refine",  # This might need adjustment based on the actual functionality available.
#         return_intermediate_steps=True,
#         question_prompt=PROMPT
#     )

# resp = chain({"input_documents": docs_4}, return_only_outputs=True)

# #add_to_document(resp["output_text"])

# resp["output_text"]

# CONTENTS = resp["output_text"]

# # 文を分割し、箇条書きで表示
# sentences = CONTENTS.split('. ')
# for sentence in sentences:
#     print(f"- {sentence.strip()}")



#HR##################################    

# prompt_template = """As a CEO of the company, listen to the report provided at the status meeting by Human resources general manager of a company.
# Company in located in Ukraine and is a official distributor of Toyota and Lexus vehicles and spare parts. 
# Please compile a minutes of the meeting report as you are a CEO and make it reopr to General director of the company based on the provided information.
# Your memo should adhere to the following guidelines:
#     - Tone: Formal
#     - Format: Detailed report divided into chapters as outlined below
#     - Length: Approximately 3000-4000 words
#     - Focus on the subject matter of discussions rather than the participants
#     - Provide detailed explanations and analysis for each topic
#     - Prepare report with bullet points for each chapter. Bullet moints are must. After each bullet poit make line break
#     - Please exlude introduction, greeting, closing part and signature. Only report

        
# Detailed insights, data comparisons, market trends, and strategic conclusions are crucial. Where applicable, include follow-up actions with specific owners assigned.

# Your analysis should incorporate the following discussions and data points:
# {text}

# Ensure precision and depth in your report:"""
# PROMPT = PromptTemplate(template=prompt_template, input_variables=["text"])

# chain = load_summarize_chain(
#         openaichat,
#         chain_type="refine",  # This might need adjustment based on the actual functionality available.
#         return_intermediate_steps=True,
#         question_prompt=PROMPT
#     )

# resp = chain({"input_documents": docs_4}, return_only_outputs=True)

# #add_to_document(resp["output_text"])

# resp["output_text"]

# CONTENTS = resp["output_text"]

# # 文を分割し、箇条書きで表示
# sentences = CONTENTS.split('. ')
# for sentence in sentences:
#     print(f"- {sentence.strip()}")




#BRAND EQUITY##################################    

# prompt_template = """As a CEO of the company, listen to the report provided at the status meeting by Customer Relationship and Dealers Development general manager of a company.
# Company in located in Ukraine and is a official distributor of Toyota and Lexus vehicles and spare parts. 
# Please compile a minutes of the meeting report as you are a CEO and make it reopr to General director of the company based on the provided information.
# Your memo should adhere to the following guidelines:
#     - Tone: Formal
#     - Format: Detailed report divided into chapters as outlined below
#     - Length: Approximately 3000-4000 words
#     - Focus on the subject matter of discussions rather than the participants
#     - Provide detailed explanations and analysis for each topic
#     - Prepare report with bullet points for each chapter. Bullet moints are must. After each bullet poit make line break
#     - Please exlude introduction, greeting, closing part and signature. Only report
#     - Organize the memo into chapters on:
#         1. Dealers network development relatedissues 
#         2. Customer relationship related issues

        
# Detailed insights, data comparisons, market trends, and strategic conclusions are crucial. Where applicable, include follow-up actions with specific owners assigned.

# Your analysis should incorporate the following discussions and data points:
# {text}

# Ensure precision and depth in your report:"""
# PROMPT = PromptTemplate(template=prompt_template, input_variables=["text"])

# chain = load_summarize_chain(
#         openaichat,
#         chain_type="refine",  # This might need adjustment based on the actual functionality available.
#         return_intermediate_steps=True,
#         question_prompt=PROMPT
#     )

# resp = chain({"input_documents": docs_4}, return_only_outputs=True)

# #add_to_document(resp["output_text"])

# resp["output_text"]

# CONTENTS = resp["output_text"]

# # 文を分割し、箇条書きで表示
# sentences = CONTENTS.split('. ')
# for sentence in sentences:
#     print(f"- {sentence.strip()}")



#FREE DISCUSIION  !!!!!!!!##################################    

# prompt_template = """As a CEO of the company, listen to the report provided at the status meeting by Customer Relationship and Dealers Development general manager of a company.
# Company in located in Ukraine and is a official distributor of Toyota and Lexus vehicles and spare parts. 
# Please compile a minutes of the meeting report as you are a CEO and make it reopr to General director of the company based on the provided information.
# Your memo should adhere to the following guidelines:
#     - Tone: Formal
#     - Format: Detailed report divided into chapters as outlined below
#     - Length: Approximately 3000-4000 words
#     - Focus on the subject matter of discussions rather than the participants
#     - Provide detailed explanations and analysis for each topic
#     - Prepare report with bullet points for each chapter. Bullet moints are must. After each bullet poit make line break
#     - Please exlude introduction, greeting, closing part and signature. Only report
#     - Organize the memo into chapters on:
#         1. Dealers network development relatedissues 
#         2. Customer relationship related issues

        
# Detailed insights, data comparisons, market trends, and strategic conclusions are crucial. Where applicable, include follow-up actions with specific owners assigned.

# Your analysis should incorporate the following discussions and data points:
# {text}

# Ensure precision and depth in your report:"""
# PROMPT = PromptTemplate(template=prompt_template, input_variables=["text"])

# chain = load_summarize_chain(
#         openaichat,
#         chain_type="refine",  # This might need adjustment based on the actual functionality available.
#         return_intermediate_steps=True,
#         question_prompt=PROMPT
#     )

# resp = chain({"input_documents": docs_4}, return_only_outputs=True)

# #add_to_document(resp["output_text"])

# resp["output_text"]

# CONTENTS = resp["output_text"]

# # 文を分割し、箇条書きで表示
# sentences = CONTENTS.split('. ')
# for sentence in sentences:
#     print(f"- {sentence.strip()}")





#SPECIFIC ISSUE  !!!!!!!!##################################    

# prompt_template = """As a CEO of the company, listen to the report provided at the status meeting by Customer Relationship and Dealers Development general manager of a company.
# Company in located in Ukraine and is a official distributor of Toyota and Lexus vehicles and spare parts. 
# Please compile a minutes of the meeting report as you are a CEO and make it reopr to General director of the company based on the provided information.
# Your memo should adhere to the following guidelines:
#     - Tone: Formal
#     - Format: Detailed report divided into chapters as outlined below
#     - Length: Approximately 3000-4000 words
#     - Focus on the subject matter of discussions rather than the participants
#     - Provide detailed explanations and analysis for each topic
#     - Prepare report with bullet points for each chapter. Bullet moints are must. After each bullet poit make line break
#     - Please exlude introduction, greeting, closing part and signature. Only report
#     - Organize the memo into chapters on:
#         1. Dealers network development relatedissues 
#         2. Customer relationship related issues

        
# Detailed insights, data comparisons, market trends, and strategic conclusions are crucial. Where applicable, include follow-up actions with specific owners assigned.

# Your analysis should incorporate the following discussions and data points:
# {text}

# Ensure precision and depth in your report:"""
# PROMPT = PromptTemplate(template=prompt_template, input_variables=["text"])

# chain = load_summarize_chain(
#         openaichat,
#         chain_type="refine",  # This might need adjustment based on the actual functionality available.
#         return_intermediate_steps=True,
#         question_prompt=PROMPT
#     )

# resp = chain({"input_documents": docs_4}, return_only_outputs=True)

# #add_to_document(resp["output_text"])

# resp["output_text"]

# CONTENTS = resp["output_text"]

# # 文を分割し、箇条書きで表示
# sentences = CONTENTS.split('. ')
# for sentence in sentences:
#     print(f"- {sentence.strip()}")




#GD  !!!!!!!!##################################    

# prompt_template = """As a CEO of the company, listen to the report provided at the status meeting by Customer Relationship and Dealers Development general manager of a company.
# Company in located in Ukraine and is a official distributor of Toyota and Lexus vehicles and spare parts. 
# Please compile a minutes of the meeting report as you are a CEO and make it reopr to General director of the company based on the provided information.
# Your memo should adhere to the following guidelines:
#     - Tone: Formal
#     - Format: Detailed report divided into chapters as outlined below
#     - Length: Approximately 3000-4000 words
#     - Focus on the subject matter of discussions rather than the participants
#     - Provide detailed explanations and analysis for each topic
#     - Prepare report with bullet points for each chapter. Bullet moints are must. After each bullet poit make line break
#     - Please exlude introduction, greeting, closing part and signature. Only report
#     - Organize the memo into chapters on:
#         1. Dealers network development relatedissues 
#         2. Customer relationship related issues

        
# Detailed insights, data comparisons, market trends, and strategic conclusions are crucial. Where applicable, include follow-up actions with specific owners assigned.

# Your analysis should incorporate the following discussions and data points:
# {text}

# Ensure precision and depth in your report:"""
# PROMPT = PromptTemplate(template=prompt_template, input_variables=["text"])

# chain = load_summarize_chain(
#         openaichat,
#         chain_type="refine",  # This might need adjustment based on the actual functionality available.
#         return_intermediate_steps=True,
#         question_prompt=PROMPT
#     )

# resp = chain({"input_documents": docs_4}, return_only_outputs=True)

# #add_to_document(resp["output_text"])

# resp["output_text"]

# CONTENTS = resp["output_text"]

# # 文を分割し、箇条書きで表示
# sentences = CONTENTS.split('. ')
# for sentence in sentences:
#     print(f"- {sentence.strip()}")


def add_text_with_title_to_doc(title, text, filename="MoM - th Toyota Ukraine Management Meeting.docx"):
    # ファイルが存在するか確認し、適切なDocumentオブジェクトを生成
    if os.path.exists(filename):
        doc = Document(filename)  # 既存のドキュメントを開く
        print(f"Opening existing document: {filename}")
    else:
        doc = Document()  # 新しいドキュメントを作成
        print(f"Creating new document: {filename}")

    # タイトルをドキュメントに追加し、太字に設定
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    
    # タイトルとテキストの間に改行を挿入
    doc.add_paragraph()
    
    # テキストデータをドキュメントに追加
    text_paragraph = doc.add_paragraph(text)

    # ドキュメントを保存
    doc.save(filename)
    print(f"Document saved as '{filename}'.")

# 例示用のデータ
agenda_title = df['Agenda'].unique()[2]  # タイトルとして使用するAgendaを取得
contents = resp["output_text"]  # APIレスポンスからテキストデータを取得

# 関数を呼び出してタイトルとテキストをWordドキュメントに追加
add_text_with_title_to_doc(agenda_title, contents)

